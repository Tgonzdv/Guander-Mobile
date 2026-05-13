# -*- coding: utf-8 -*-
"""
Genera el Manual de Usuario de Guander en formato .docx
Estilo similar al Manual de Fast Track
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── COLORES ──────────────────────────────────────────────────────────────────
COLOR_PRIMARIO   = RGBColor(0x23, 0x96, 0xF6)   # Azul Guander
COLOR_TITULO     = RGBColor(0x1A, 0x1A, 0x2E)   # Casi negro
COLOR_GRIS       = RGBColor(0x55, 0x55, 0x55)
COLOR_BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_FONDO_TBL  = RGBColor(0x23, 0x96, 0xF6)
COLOR_FONDO_NOTA = RGBColor(0xE8, 0xF4, 0xFD)   # Azul muy claro

# ─── HELPERS ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Pinta el fondo de una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_run_color(run, color: RGBColor):
    run.font.color.rgb = color


def add_heading(doc, text, level=1, color=None):
    """Agrega un heading con color opcional."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if color:
            run.font.color.rgb = color
    return p


def add_body(doc, text, bold=False, color=None, indent=False):
    """Agrega un párrafo de cuerpo."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_nota(doc, texto_bold, texto_normal="", tipo="IMPORTANTE"):
    """Cuadro IMPORTANTE / ATENCIÓN."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{tipo}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)
    if texto_bold:
        r2 = p.add_run(texto_bold)
        r2.bold = True
        r2.font.size = Pt(11)
    if texto_normal:
        r3 = p.add_run(texto_normal)
        r3.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    """Agrega un ítem de lista con bullet."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_numbered_item(doc, text, bold_prefix=""):
    """Agrega un ítem de lista numerada."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_step(doc, number, text):
    """Agrega un paso numerado."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    rb = p.add_run(f"{number}. ")
    rb.bold = True
    rb.font.size = Pt(11)
    rb.font.color.rgb = COLOR_PRIMARIO
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(9)
    return p


def table_header_row(table, headers, hex_color="2396F6"):
    """Agrega la primera fila como encabezado con fondo de color."""
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, hex_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = COLOR_BLANCO
        run.font.size = Pt(10)


def fill_table_row(table, row_idx, values, center_cols=None):
    center_cols = center_cols or []
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        if i in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.size = Pt(10)


# ─── PORTADA ──────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Espaciado superior
    for _ in range(6):
        doc.add_paragraph()

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Guander")
    r.bold = True
    r.font.size = Pt(48)
    r.font.color.rgb = COLOR_PRIMARIO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Manual de Usuario")
    r2.bold = True
    r2.font.size = Pt(28)
    r2.font.color.rgb = COLOR_TITULO

    for _ in range(2):
        doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Tu compañero pet-friendly en cada salida")
    r3.italic = True
    r3.font.size = Pt(14)
    r3.font.color.rgb = COLOR_GRIS

    for _ in range(6):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Versión: 1.0")
    r4.font.size = Pt(12)
    r4.font.color.rgb = COLOR_GRIS

    doc.add_page_break()


# ─── ÍNDICE ────────────────────────────────────────────────────────────────────

def build_toc(doc):
    add_heading(doc, "Índice", level=1, color=COLOR_PRIMARIO)

    toc_items = [
        ("1.", "Introducción", "3"),
        ("2.", "Objetivo", "3"),
        ("3.", "Requisitos del sistema", "4"),
        ("4.", "Instalación de la app", "4"),
        ("5.", "Ingreso al sistema", "5"),
        ("   5.1.", "Iniciar sesión con Google", "5"),
        ("   5.2.", "Registrarse con email y contraseña", "5"),
        ("   5.3.", "Iniciar sesión con email y contraseña", "6"),
        ("   5.4.", "Cerrar sesión", "6"),
        ("6.", "Panel Principal (Dashboard)", "7"),
        ("   6.1.", "Vista del Dashboard", "7"),
        ("   6.2.", "Barra de navegación inferior", "7"),
        ("7.", "Mapa y Exploración de Locales", "8"),
        ("   7.1.", "Ver el mapa interactivo", "8"),
        ("   7.2.", "Lista de locales cercanos", "8"),
        ("   7.3.", "Filtrar por categoría", "9"),
        ("   7.4.", "Buscar un local", "9"),
        ("   7.5.", "Ver detalle de un local", "9"),
        ("   7.6.", "Dejar una reseña", "10"),
        ("8.", "Escaneo de Códigos QR y PetPoints", "11"),
        ("   8.1.", "Cómo funciona el sistema de puntos", "11"),
        ("   8.2.", "Escanear un código QR", "11"),
        ("   8.3.", "Resultado del escaneo", "12"),
        ("9.", "Recompensas", "13"),
        ("   9.1.", "Catálogo de recompensas", "13"),
        ("   9.2.", "Buscar recompensas", "13"),
        ("   9.3.", "Canjear una recompensa", "14"),
        ("   9.4.", "Historial de canjes", "14"),
        ("10.", "Perfil de Usuario", "15"),
        ("   10.1.", "Ver mi perfil", "15"),
        ("   10.2.", "Editar datos personales", "15"),
        ("   10.3.", "Cambiar foto de perfil", "16"),
        ("11.", "Configuración", "17"),
        ("   11.1.", "Notificaciones", "17"),
        ("   11.2.", "Privacidad y seguridad", "17"),
        ("   11.3.", "Apariencia (tema claro / oscuro)", "18"),
        ("   11.4.", "Idioma", "18"),
        ("   11.5.", "Centro de Ayuda", "18"),
        ("   11.6.", "Términos y Condiciones", "18"),
        ("   11.7.", "Política de Privacidad", "19"),
        ("12.", "Preguntas Frecuentes (FAQ)", "20"),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if num.strip().count(".") == 1 and not num.startswith(" "):
            # Línea de sección principal
            rn = p.add_run(f"{num} {title}")
            rn.bold = True
            rn.font.size = Pt(11)
            rn.font.color.rgb = COLOR_TITULO
        else:
            rn = p.add_run(f"{num} {title}")
            rn.font.size = Pt(10.5)
            rn.font.color.rgb = COLOR_GRIS

    doc.add_page_break()


# ─── SECCIÓN 1 — INTRODUCCIÓN ─────────────────────────────────────────────────

def sec1_introduccion(doc):
    add_heading(doc, "1. Introducción", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "Bienvenido/a a Guander, la app Android diseñada especialmente para dueños de mascotas "
        "que quieren disfrutar al máximo los establecimientos y profesionales pet-friendly de su ciudad.")
    doc.add_paragraph()
    add_body(doc,
        "Con Guander podés explorar locales adheridos en un mapa interactivo, acumular "
        "PetPoints escaneando los códigos QR de cada visita y canjearlos por recompensas "
        "exclusivas. También podés ver reseñas de otros usuarios y dejar la tuya propia.")
    doc.add_paragraph()
    add_body(doc,
        "La plataforma fue construida sobre Android nativo (Java) con Material Design 3, "
        "mapas OpenStreetMap, escáner ZXing y un backend en Cloudflare Workers con base de "
        "datos D1 (SQLite serverless) y autenticación Firebase.")
    doc.add_paragraph()


# ─── SECCIÓN 2 — OBJETIVO ─────────────────────────────────────────────────────

def sec2_objetivo(doc):
    add_heading(doc, "2. Objetivo", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "Si bien Guander fue desarrollada bajo un riguroso estándar de experiencia de usuario "
        "—haciéndola intuitiva y autodidacta—, a través de este manual podrás conocer en "
        "detalle cada funcionalidad, cómo utilizarla y familiarizarte con los flujos de la app.")
    doc.add_paragraph()
    add_body(doc,
        "Este documento se encuentra en desarrollo continuo y acompaña cada nueva versión "
        "publicada en Google Play. Cualquier cambio relevante será comunicado a través de las "
        "notas de actualización de la tienda.")
    doc.add_paragraph()
    add_body(doc,
        "Esperamos que sea de utilidad. Estamos a disposición para cualquier consulta o "
        "sugerencia.")
    doc.add_paragraph()


# ─── SECCIÓN 3 — REQUISITOS ───────────────────────────────────────────────────

def sec3_requisitos(doc):
    add_heading(doc, "3. Requisitos del sistema", level=1, color=COLOR_PRIMARIO)

    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = 'Table Grid'
    table_header_row(tbl, ["Requisito", "Detalle"])
    rows_data = [
        ("Sistema operativo", "Android 7.0 (API 24) o superior"),
        ("Versión recomendada", "Android 10 o superior"),
        ("Espacio en disco", "~35 MB de almacenamiento libre"),
        ("Conexión a Internet", "Requerida para la mayoría de las funciones"),
        ("Cámara", "Necesaria para escanear códigos QR"),
    ]
    for i, (a, b) in enumerate(rows_data, start=1):
        fill_table_row(tbl, i, [a, b])

    doc.add_paragraph()
    add_nota(doc, "La app no es compatible con iOS.",
             " Guander es una aplicación exclusiva para dispositivos Android.")
    doc.add_paragraph()


# ─── SECCIÓN 4 — INSTALACIÓN ──────────────────────────────────────────────────

def sec4_instalacion(doc):
    add_heading(doc, "4. Instalación de la app", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "La instalación de Guander se realiza directamente desde la tienda oficial de "
        "aplicaciones de Android.")
    doc.add_paragraph()

    add_step(doc, 1, "Abrí Google Play Store en tu dispositivo Android.")
    add_step(doc, 2, 'En el buscador escribí "Guander".')
    add_step(doc, 3, 'Seleccioná la app y tocá el botón "Instalar".')
    add_step(doc, 4, "Esperá a que finalice la descarga e instalación.")
    add_step(doc, 5, 'Abrí la app tocando "Abrir" o el ícono de Guander en tu pantalla de inicio.')

    doc.add_paragraph()
    add_nota(doc, "",
             "También podés instalar la app escaneando el código QR disponible en el sitio "
             "oficial de Guander.")
    doc.add_paragraph()

    doc.add_page_break()


# ─── SECCIÓN 5 — INGRESO AL SISTEMA ──────────────────────────────────────────

def sec5_ingreso(doc):
    add_heading(doc, "5. Ingreso al sistema", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "Al abrir Guander por primera vez, serás dirigido a la pantalla de inicio de sesión. "
        "La app ofrece dos métodos de autenticación.")
    doc.add_paragraph()

    # 5.1
    add_heading(doc, "5.1. Iniciar sesión con Google", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Este es el método más rápido. Usá tu cuenta de Google para ingresar sin tener que "
        "recordar una contraseña adicional.")
    doc.add_paragraph()
    add_step(doc, 1, 'En la pantalla de login, tocá el botón "Continuar con Google".')
    add_step(doc, 2, "Seleccioná la cuenta de Google que querés usar.")
    add_step(doc, 3, "Guander verificará si ya tenés cuenta registrada:")
    add_bullet(doc, "Si ya existe: ingresarás directamente al Dashboard.", level=1)
    add_bullet(doc, "Si es nueva: se creará tu cuenta automáticamente y recibirás un email de bienvenida.", level=1)
    doc.add_paragraph()
    add_nota(doc, "Al registrarte por primera vez,", " se te pedirá ingresar tu nombre para completar el perfil.")
    doc.add_paragraph()

    # 5.2
    add_heading(doc, "5.2. Registrarse con email y contraseña", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Si preferís no usar tu cuenta de Google, podés crear una cuenta directamente en Guander.")
    doc.add_paragraph()
    add_step(doc, 1, 'Tocá el enlace "Crear cuenta" en la pantalla de login.')
    add_step(doc, 2, "Completá los siguientes campos:")
    add_bullet(doc, "Nombre", level=1)
    add_bullet(doc, "Apellido", level=1)
    add_bullet(doc, "Teléfono", level=1)
    add_bullet(doc, "Correo electrónico", level=1)
    add_bullet(doc, "Contraseña (mínimo 6 caracteres)", level=1)
    add_step(doc, 3, 'Tocá "Registrarme".')
    add_step(doc, 4, "Recibirás un email de bienvenida y serás dirigido al Dashboard.")
    doc.add_paragraph()

    # 5.3
    add_heading(doc, "5.3. Iniciar sesión con email y contraseña", level=2, color=COLOR_TITULO)
    add_body(doc, "Si ya tenés cuenta registrada con email y contraseña:")
    doc.add_paragraph()
    add_step(doc, 1, "Ingresá tu correo electrónico y contraseña en los campos correspondientes.")
    add_step(doc, 2, 'Tocá "Ingresar".')
    add_step(doc, 3, "Serás redirigido al Dashboard si las credenciales son correctas.")
    doc.add_paragraph()
    add_nota(doc, "En caso de credenciales incorrectas,",
             " el sistema mostrará un mensaje de error en español indicando el problema.")
    doc.add_paragraph()

    # 5.4
    add_heading(doc, "5.4. Cerrar sesión", level=2, color=COLOR_TITULO)
    add_body(doc,
        "La sesión permanece activa entre reinicios de la app para una experiencia más cómoda. "
        "Para cerrar sesión manualmente:")
    doc.add_paragraph()
    add_step(doc, 1, "Ingresá a la sección Perfil desde la barra de navegación inferior.")
    add_step(doc, 2, 'Tocá el botón "Cerrar sesión" en la parte inferior de la pantalla.')
    add_step(doc, 3, "La app cerrará la sesión de Firebase y borrará los datos locales.")
    doc.add_paragraph()
    add_nota(doc, "Al cerrar sesión",
             " tanto la autenticación de Google como los datos de sesión local son eliminados "
             "del dispositivo.")
    doc.add_paragraph()

    doc.add_page_break()


# ─── SECCIÓN 6 — DASHBOARD ────────────────────────────────────────────────────

def sec6_dashboard(doc):
    add_heading(doc, "6. Panel Principal (Dashboard)", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "El Dashboard es la pantalla principal de Guander. Al ingresar, encontrarás un resumen "
        "de tu actividad y accesos rápidos a todas las secciones.")
    doc.add_paragraph()

    # 6.1
    add_heading(doc, "6.1. Vista del Dashboard", level=2, color=COLOR_TITULO)
    add_body(doc, "La pantalla principal contiene los siguientes elementos:")
    doc.add_paragraph()

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    table_header_row(tbl, ["Elemento", "Descripción"])
    rows_data = [
        ("Saludo personalizado", "Muestra tu nombre y un saludo de bienvenida."),
        ("Saldo de PetPoints", "Indica tu saldo actual de puntos acumulados con separadores de miles."),
        ("Últimas notificaciones", "Muestra las últimas 5 notificaciones con ícono, título y descripción."),
        ("Accesos rápidos", "Botones de acceso directo a Mapa, Recompensas, QR y Perfil."),
    ]
    for i, (a, b) in enumerate(rows_data, start=1):
        fill_table_row(tbl, i, [a, b])

    doc.add_paragraph()

    # 6.2
    add_heading(doc, "6.2. Barra de navegación inferior", level=2, color=COLOR_TITULO)
    add_body(doc,
        "En la parte inferior de la pantalla encontrarás la barra de navegación principal, "
        "disponible en todo momento en la app:")
    doc.add_paragraph()

    tbl2 = doc.add_table(rows=6, cols=3)
    tbl2.style = 'Table Grid'
    table_header_row(tbl2, ["Ícono / Sección", "Nombre", "Función"])
    rows2 = [
        ("🏠  Inicio", "Dashboard", "Volver al panel principal con tus puntos y notificaciones."),
        ("🗺️  Mapa", "Mapa", "Explorar locales pet-friendly cercanos en el mapa."),
        ("📷  QR", "Escáner QR", "Escanear el código QR de un local para sumar PetPoints."),
        ("🎁  Puntos", "Recompensas", "Ver el catálogo de recompensas y canjear puntos."),
        ("👤  Perfil", "Perfil", "Ver y editar tu perfil, ajustes y configuración."),
    ]
    for i, row in enumerate(rows2, start=1):
        fill_table_row(tbl2, i, row)

    doc.add_paragraph()
    doc.add_page_break()


# ─── SECCIÓN 7 — MAPA ─────────────────────────────────────────────────────────

def sec7_mapa(doc):
    add_heading(doc, "7. Mapa y Exploración de Locales", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "La sección Mapa es el corazón de Guander. Desde aquí podés descubrir todos los "
        "establecimientos y profesionales pet-friendly cercanos a tu ubicación.")
    doc.add_paragraph()

    # 7.1
    add_heading(doc, "7.1. Ver el mapa interactivo", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Al ingresar a la sección Mapa, la app solicitará permiso para acceder a tu ubicación. "
        "El mapa se centrará automáticamente en tu posición actual.")
    doc.add_paragraph()
    add_body(doc,
        "Cada establecimiento adherido aparece como un marcador en el mapa. Podés pellizcar "
        "para hacer zoom, deslizar para navegar y tocar un marcador para ver el nombre del local.")
    doc.add_paragraph()
    add_nota(doc, "Si no otorgás permiso de ubicación,",
             " el mapa se centrará en Buenos Aires como ubicación predeterminada.")
    doc.add_paragraph()

    # 7.2
    add_heading(doc, "7.2. Lista de locales cercanos", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Debajo del mapa se muestra una lista de todos los establecimientos, ordenados por "
        "distancia de menor a mayor. Cada tarjeta contiene:")
    doc.add_paragraph()
    add_bullet(doc, "Nombre del establecimiento")
    add_bullet(doc, "Distancia desde tu ubicación (en metros o kilómetros)")
    add_bullet(doc, "Calificación promedio con estrellas")
    add_bullet(doc, "Horarios (días de semana, fin de semana y domingo)")
    add_bullet(doc, "Estado actual: Abierto / Cerrado")
    doc.add_paragraph()

    # 7.3
    add_heading(doc, "7.3. Filtrar por categoría", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Utilizá los botones de filtro en la parte superior de la lista para ver únicamente "
        "los establecimientos de una categoría específica:")
    doc.add_paragraph()

    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = 'Table Grid'
    table_header_row(tbl, ["Categoría", "Tipo de establecimiento"])
    cats = [
        ("Todos", "Muestra todos los establecimientos sin filtro."),
        ("Locales", "Tiendas y negocios pet-friendly."),
        ("Restaurantes", "Bares, cafés y restaurantes que admiten mascotas."),
        ("Profesionales", "Veterinarios, peluqueros caninos, adiestradores, etc."),
        ("Servicios", "Guarderías, paseos caninos y otros servicios."),
    ]
    for i, (a, b) in enumerate(cats, start=1):
        fill_table_row(tbl, i, [a, b])

    doc.add_paragraph()
    add_body(doc,
        "Al seleccionar un filtro, los marcadores del mapa y la lista se actualizan automáticamente.")
    doc.add_paragraph()

    # 7.4
    add_heading(doc, "7.4. Buscar un local", level=2, color=COLOR_TITULO)
    add_body(doc,
        "La barra de búsqueda en la parte superior permite buscar establecimientos por nombre "
        "o descripción en tiempo real. A medida que escribís, la lista y el mapa se actualizan "
        "con los resultados coincidentes.")
    doc.add_paragraph()

    # 7.5
    add_heading(doc, "7.5. Ver detalle de un local", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Al tocar una tarjeta de la lista o un marcador del mapa, accedés a la pantalla de "
        "detalle del establecimiento. Esta pantalla contiene:")
    doc.add_paragraph()
    add_bullet(doc, "Foto del establecimiento")
    add_bullet(doc, "Nombre y categoría (con ícono)")
    add_bullet(doc, "Estado: Abierto / Cerrado")
    add_bullet(doc, "Descripción completa")
    add_bullet(doc, "Dirección y teléfono de contacto")
    add_bullet(doc, "Botón para abrir la dirección en Google Maps o app de navegación")
    add_bullet(doc, "Lista completa de reseñas de usuarios")
    doc.add_paragraph()

    # 7.6
    add_heading(doc, "7.6. Dejar una reseña", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Si visitaste un local y realizaste una compra escaneando su QR, podés dejar tu reseña:")
    doc.add_paragraph()
    add_step(doc, 1, "Entrá al detalle del establecimiento.")
    add_step(doc, 2, 'Tocá el botón "Dejar reseña" o deslizá hasta la sección de comentarios.')
    add_step(doc, 3, "Seleccioná tu puntuación de 1 a 5 estrellas.")
    add_step(doc, 4, "Escribí tu comentario.")
    add_step(doc, 5, 'Tocá "Enviar".')
    doc.add_paragraph()
    add_nota(doc, "Solo podés dejar una reseña",
             " por establecimiento, y únicamente si tenés al menos una compra registrada allí. "
             "Si aún no visitaste el local, verás el mensaje \"No podés comentar aún\".")
    doc.add_paragraph()

    doc.add_page_break()


# ─── SECCIÓN 8 — QR ──────────────────────────────────────────────────────────

def sec8_qr(doc):
    add_heading(doc, "8. Escaneo de Códigos QR y PetPoints", level=1, color=COLOR_PRIMARIO)

    # 8.1
    add_heading(doc, "8.1. Cómo funciona el sistema de puntos", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Cada establecimiento adherido a Guander cuenta con un código QR propio. "
        "Cuando realizás una compra o visita, el local te muestra su QR para que lo escanees "
        "con la app. El sistema convierte el monto de la operación en PetPoints según la "
        "siguiente fórmula:")
    doc.add_paragraph()

    p_formula = doc.add_paragraph()
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_formula = p_formula.add_run("PetPoints = máx(1, parte entera de (monto / 1000))")
    r_formula.bold = True
    r_formula.font.size = Pt(12)
    r_formula.font.color.rgb = COLOR_PRIMARIO

    doc.add_paragraph()
    add_body(doc, "Ejemplos:")
    add_bullet(doc, "Compra de $2.500 → 2 PetPoints")
    add_bullet(doc, "Compra de $800 → 1 PetPoint (mínimo garantizado)")
    add_bullet(doc, "Compra de $10.000 → 10 PetPoints")
    doc.add_paragraph()
    add_body(doc,
        "Acumulá PetPoints en cada visita y utilizalos para canjear recompensas exclusivas "
        "en la sección Recompensas.")
    doc.add_paragraph()

    # 8.2
    add_heading(doc, "8.2. Escanear un código QR", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Para sumar PetPoints desde la app, seguí estos pasos:")
    doc.add_paragraph()
    add_step(doc, 1, 'Tocá el ícono de QR (📷) en la barra de navegación inferior.')
    add_step(doc, 2,
        "La app solicitará permiso para usar la cámara. Aceptá para continuar.")
    add_step(doc, 3,
        "Enfocá la cámara sobre el código QR que te muestra el establecimiento.")
    add_step(doc, 4,
        "La app detecta el QR automáticamente en modo one-shot (un escaneo a la vez).")
    doc.add_paragraph()
    add_nota(doc, "Si el escaneo falla",
             " o necesitás volver a intentarlo, tocá el botón de reinicio para activar "
             "nuevamente el escáner.")
    doc.add_paragraph()

    # 8.3
    add_heading(doc, "8.3. Resultado del escaneo", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Si el código QR es válido, la app mostrará un diálogo de confirmación con:")
    doc.add_paragraph()
    add_bullet(doc, "Puntos ganados en esta visita")
    add_bullet(doc, "Nombre del local visitado")
    add_bullet(doc, "Descripción del ítem / servicio adquirido")
    add_bullet(doc, "Monto de la operación")
    add_bullet(doc, "Nuevo saldo total de PetPoints")
    doc.add_paragraph()
    add_nota(doc, "QR inválido:",
             " si el código no pertenece a Guander o fue alterado, el sistema mostrará "
             "un mensaje de error y no acreditará puntos.")
    doc.add_paragraph()

    doc.add_page_break()


# ─── SECCIÓN 9 — RECOMPENSAS ─────────────────────────────────────────────────

def sec9_recompensas(doc):
    add_heading(doc, "9. Recompensas", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "La sección Recompensas te permite ver qué podés obtener con tus PetPoints acumulados "
        "y canjearlos por beneficios exclusivos en los establecimientos adheridos.")
    doc.add_paragraph()

    # 9.1
    add_heading(doc, "9.1. Catálogo de recompensas", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Al ingresar a la sección Recompensas, verás el catálogo completo de beneficios "
        "disponibles. Cada tarjeta de recompensa muestra:")
    doc.add_paragraph()
    add_bullet(doc, "Nombre de la recompensa")
    add_bullet(doc, "Descripción del beneficio")
    add_bullet(doc, "Costo en PetPoints")
    add_bullet(doc, "Ícono o logo del local que lo ofrece")
    add_bullet(doc, "Botón de canje (activo si tenés suficientes puntos, deshabilitado si no)")
    doc.add_paragraph()
    add_nota(doc, "Las recompensas ya vencidas o que ya canjeaste",
             " no aparecerán en el catálogo.")
    doc.add_paragraph()

    # 9.2
    add_heading(doc, "9.2. Buscar recompensas", level=2, color=COLOR_TITULO)
    add_body(doc,
        "La barra de búsqueda en la parte superior del catálogo filtra las recompensas por "
        "nombre o descripción en tiempo real. Escribí el nombre del local o la recompensa "
        "para encontrarla rápidamente.")
    doc.add_paragraph()

    # 9.3
    add_heading(doc, "9.3. Canjear una recompensa", level=2, color=COLOR_TITULO)
    add_body(doc, "Para canjear una recompensa, seguí estos pasos:")
    doc.add_paragraph()
    add_step(doc, 1, "Asegurate de tener suficientes PetPoints (el botón estará activo).")
    add_step(doc, 2, 'Tocá "Canjear" en la recompensa deseada.')
    add_step(doc, 3,
        "Un diálogo de confirmación mostrará el nombre de la recompensa y el costo en puntos.")
    add_step(doc, 4, 'Tocá "Confirmar" para proceder.')
    add_step(doc, 5,
        "El sistema deducirá los puntos de tu saldo y te mostrará el código de canje.")
    add_step(doc, 6,
        "Presentá el código de canje en el establecimiento para hacer efectivo el beneficio.")
    doc.add_paragraph()
    add_nota(doc, "Cada recompensa solo puede canjearse una vez.",
             " Si intentás canjear una recompensa que ya utilizaste, el sistema lo impedirá.")
    doc.add_paragraph()

    # 9.4
    add_heading(doc, "9.4. Historial de canjes", level=2, color=COLOR_TITULO)
    add_body(doc,
        "En la pestaña \"Historial\" de la sección Recompensas podés ver todos tus canjes "
        "anteriores. Cada entrada muestra:")
    doc.add_paragraph()
    add_bullet(doc, "Ícono del beneficio")
    add_bullet(doc, "Descripción de la recompensa")
    add_bullet(doc, "Fecha relativa del canje (Hoy / Hace N días)")
    add_bullet(doc, "PetPoints utilizados")
    add_bullet(doc, "Código de canje")
    doc.add_paragraph()

    doc.add_page_break()


# ─── SECCIÓN 10 — PERFIL ─────────────────────────────────────────────────────

def sec10_perfil(doc):
    add_heading(doc, "10. Perfil de Usuario", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "La sección Perfil te permite ver y gestionar tu información personal dentro de Guander.")
    doc.add_paragraph()

    # 10.1
    add_heading(doc, "10.1. Ver mi perfil", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Al ingresar al Perfil verás un resumen completo de tu cuenta:")
    doc.add_paragraph()

    tbl = doc.add_table(rows=8, cols=2)
    tbl.style = 'Table Grid'
    table_header_row(tbl, ["Campo", "Descripción"])
    rows_data = [
        ("Avatar", "Tu foto de perfil o la inicial de tu nombre en un círculo de color si no tenés foto."),
        ("Nombre completo", "Nombre y apellido registrados."),
        ("Email", "Dirección de correo electrónico asociada a la cuenta."),
        ("Teléfono", "Número de contacto."),
        ("Dirección", "Domicilio ingresado en el perfil."),
        ("Fecha de alta", "Fecha en que creaste tu cuenta en Guander."),
        ("Estadísticas", "Lugares visitados, PetPoints acumulados y cupones canjeados."),
    ]
    for i, (a, b) in enumerate(rows_data, start=1):
        fill_table_row(tbl, i, [a, b])

    doc.add_paragraph()

    # 10.2
    add_heading(doc, "10.2. Editar datos personales", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Para actualizar tu información de perfil:")
    doc.add_paragraph()
    add_step(doc, 1, 'Tocá el botón "Editar perfil" en la pantalla de Perfil.')
    add_step(doc, 2, "Modificá los campos que desees: nombre, apellido, teléfono o dirección.")
    add_step(doc, 3, 'Tocá "Guardar" para confirmar los cambios.')
    add_step(doc, 4, "La app actualizará tus datos y mostrará un mensaje de confirmación.")
    doc.add_paragraph()

    # 10.3
    add_heading(doc, "10.3. Cambiar foto de perfil", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Podés personalizar tu perfil con una foto propia:")
    doc.add_paragraph()
    add_step(doc, 1, "En la pantalla de Perfil, tocá tu avatar o foto actual.")
    add_step(doc, 2, "La app abrirá el selector de imágenes de tu galería.")
    add_step(doc, 3, "Seleccioná la imagen que querés usar como foto de perfil.")
    add_step(doc, 4,
        "Guander subirá la imagen automáticamente. Verás un indicador de progreso "
        "mientras se procesa.")
    add_step(doc, 5, "Una vez completada la subida, tu nueva foto aparecerá en el perfil.")
    doc.add_paragraph()
    add_nota(doc, "La foto de perfil",
             " se almacena de forma segura en la nube (Cloudflare R2) y es accesible "
             "desde cualquier dispositivo donde uses tu cuenta.")
    doc.add_paragraph()

    doc.add_page_break()


# ─── SECCIÓN 11 — CONFIGURACIÓN ──────────────────────────────────────────────

def sec11_configuracion(doc):
    add_heading(doc, "11. Configuración", level=1, color=COLOR_PRIMARIO)
    add_body(doc,
        "Desde la sección Perfil podés acceder a todas las opciones de configuración de la app. "
        "Desplazate hacia abajo para ver el menú de ajustes.")
    doc.add_paragraph()

    # 11.1
    add_heading(doc, "11.1. Notificaciones", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Controlá qué tipo de notificaciones deseás recibir. Contás con un interruptor "
        "general (master) y controles individuales por categoría:")
    doc.add_paragraph()
    add_bullet(doc, "Notificaciones generales (activa o desactiva todas las notificaciones)")
    add_bullet(doc, "Notificaciones de puntos (avisos cuando sumás PetPoints)")
    add_bullet(doc, "Notificaciones de cupones (alertas sobre tus canjes)")
    add_bullet(doc, "Lugares nuevos (cuando se suma un nuevo local a la red Guander)")
    add_bullet(doc, "Promociones (ofertas y recompensas especiales)")
    doc.add_paragraph()
    add_nota(doc, "Si desactivás las notificaciones generales,",
             " todos los demás interruptores se desactivarán automáticamente.")
    doc.add_paragraph()

    # 11.2
    add_heading(doc, "11.2. Privacidad y seguridad", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Gestioná tus preferencias de privacidad. Estas opciones son locales al dispositivo "
        "y no afectan el funcionamiento general de la app:")
    doc.add_paragraph()
    add_bullet(doc, "Compartir ubicación: permite que la app acceda a tu GPS para calcular distancias.")
    add_bullet(doc, "Datos de uso: ayuda a mejorar la app compartiendo estadísticas anónimas.")
    add_bullet(doc, "Contenido personalizado: adapta las sugerencias de locales y recompensas a tus preferencias.")
    doc.add_paragraph()

    # 11.3
    add_heading(doc, "11.3. Apariencia (tema claro / oscuro)", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Guander soporta tres modos de apariencia visual:")
    doc.add_paragraph()
    add_bullet(doc, "Sistema: sigue automáticamente la configuración de tu dispositivo.")
    add_bullet(doc, "Claro: interfaz con fondo blanco y colores brillantes.")
    add_bullet(doc, "Oscuro: interfaz con fondo oscuro, ideal para usar de noche.")
    doc.add_paragraph()
    add_body(doc,
        "Tu preferencia se guarda entre sesiones y se aplica al siguiente inicio de la app.")
    doc.add_paragraph()

    # 11.4
    add_heading(doc, "11.4. Idioma", level=2, color=COLOR_TITULO)
    add_body(doc,
        "La app se encuentra en español por defecto. Desde esta sección podés cambiar el "
        "idioma de la interfaz. La preferencia seleccionada persistirá en futuras sesiones.")
    doc.add_paragraph()

    # 11.5
    add_heading(doc, "11.5. Centro de Ayuda", level=2, color=COLOR_TITULO)
    add_body(doc,
        "El Centro de Ayuda contiene 7 preguntas frecuentes (FAQ) sobre el uso de la app, "
        "el sistema de puntos, el escaneo de QR y los canjes. Accedé desde el menú de "
        "Configuración en el Perfil.")
    doc.add_paragraph()

    # 11.6
    add_heading(doc, "11.6. Términos y Condiciones", level=2, color=COLOR_TITULO)
    add_body(doc,
        "Podés consultar los Términos y Condiciones de uso de Guander en cualquier momento "
        "desde el menú de Configuración. El documento está dividido en 8 secciones que "
        "cubren el uso del servicio, la propiedad intelectual y las responsabilidades.")
    doc.add_paragraph()

    # 11.7
    add_heading(doc, "11.7. Política de Privacidad", level=2, color=COLOR_TITULO)
    add_body(doc,
        "La Política de Privacidad detalla cómo Guander recopila, usa y protege tu información "
        "personal. Está disponible en el menú de Configuración y se divide en 8 secciones.")
    doc.add_paragraph()

    doc.add_page_break()


# ─── SECCIÓN 12 — FAQ ────────────────────────────────────────────────────────

def sec12_faq(doc):
    add_heading(doc, "12. Preguntas Frecuentes (FAQ)", level=1, color=COLOR_PRIMARIO)
    doc.add_paragraph()

    faq = [
        (
            "¿Qué son los PetPoints?",
            "Son los puntos de fidelización de Guander. Los acumulás escaneando el QR de los "
            "locales adheridos cuando realizás una compra o visita. Con ellos podés canjear "
            "recompensas y beneficios exclusivos."
        ),
        (
            "¿Cómo acumulo PetPoints?",
            "Visitá un local adherido a Guander, realizá una compra y pedile al encargado "
            "que te muestre el código QR. Escanealo con la app y los puntos se acreditarán "
            "automáticamente en tu cuenta."
        ),
        (
            "¿Por qué el botón de canje está deshabilitado?",
            "El botón de canje aparece inactivo cuando no tenés suficientes PetPoints para "
            "esa recompensa. Continuá visitando locales para acumular más puntos."
        ),
        (
            "¿Puedo dejar más de una reseña en el mismo local?",
            "No. Guander permite solo una reseña por establecimiento y por usuario. "
            "Además, para poder comentar necesitás tener al menos una compra registrada allí."
        ),
        (
            "¿Mi foto de perfil es pública?",
            "Tu foto de perfil solo es visible para vos en la pantalla de Perfil. "
            "Guander la almacena en forma segura y no la comparte con terceros."
        ),
        (
            "¿Qué hago si el escaneo de QR no funciona?",
            "Asegurate de tener buena iluminación y que el código QR esté completo en "
            "el encuadre. Si el problema persiste, tocá el botón de reinicio del escáner "
            "y volvé a intentarlo."
        ),
        (
            "¿La app funciona sin conexión a Internet?",
            "La mayoría de las funciones requieren conexión para comunicarse con el servidor. "
            "Sin embargo, las preferencias de notificaciones y apariencia se guardan "
            "localmente y están disponibles sin conexión."
        ),
    ]

    for i, (pregunta, respuesta) in enumerate(faq, start=1):
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)
        r_num = p_q.add_run(f"{i}. ")
        r_num.bold = True
        r_num.font.size = Pt(11)
        r_num.font.color.rgb = COLOR_PRIMARIO
        r_q = p_q.add_run(pregunta)
        r_q.bold = True
        r_q.font.size = Pt(11)

        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Cm(0.8)
        p_a.paragraph_format.space_after = Pt(4)
        r_a = p_a.add_run(respuesta)
        r_a.font.size = Pt(11)

    doc.add_paragraph()


# ─── FOOTER ───────────────────────────────────────────────────────────────────

def add_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guander — Manual de Usuario  |  Versión 1.0  |  Mayo 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_GRIS


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Fuente base
    style = doc.styles['Normal']
    font  = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    add_footer(doc)
    build_cover(doc)
    build_toc(doc)
    sec1_introduccion(doc)
    sec2_objetivo(doc)
    sec3_requisitos(doc)
    sec4_instalacion(doc)
    sec5_ingreso(doc)
    sec6_dashboard(doc)
    sec7_mapa(doc)
    sec8_qr(doc)
    sec9_recompensas(doc)
    sec10_perfil(doc)
    sec11_configuracion(doc)
    sec12_faq(doc)

    out_path = r"C:\Users\tomas\AndroidStudioProjects\Guander\Manual guia\Manual de usuario Guander.docx"
    doc.save(out_path)
    print(f"✓ Manual guardado en: {out_path}")


if __name__ == "__main__":
    main()
